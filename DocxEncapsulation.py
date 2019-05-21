# encoding:utf-8
from docx import Document
import os
import re


class DocxEncapsulation:
    """
    封装了python-docx的一些功能
    """

    def __init__(self, document_path: str):
        if document_path:
            if os.path.isfile(document_path) and document_path.rsplit('.', 1)[1] == 'docx':
                self.document = Document(document_path)         # 文档对象,这里是打开一个docx文档
            else:
                raise Exception("请输入正确的文件路径")
        else:
            self.document = Document()                          # 文档对象,这里是创建了一个新的文档

    def get_document(self):
        """ 获取文档对象"""
        if self.document:
            return self.document

    def get_paragraphs(self):
        """获取文档中的所有段落，文档中每一个换行符表示一个段落，这是一个list集合"""
        return self.document.paragraphs

    def matching_paragraphs(self, pattern: str):
        """
        从文档的所有段落中匹配符合正则表达式的段落，并返回相应索引，该索引表示文档中第几个段落
        :param pattern: 正则表达式
        :return: list
        """
        result_list = []
        # 遍历文档中所有段落,docx文档中每一个换行代表一个段落
        for i in range(len(self.document.paragraphs)):
            # print(i)
            if self.document.paragraphs[i].text:
                if re.match(pattern, self.document.paragraphs[i].text):
                    result_list.append(i)
        return result_list


if __name__ == '__main__':
    path = "C:\\Users\wbin3\Desktop\游戏下载链接\所有游戏下载地址.docx"
    document = DocxEncapsulation(path)
    lists = document.matching_paragraphs("^.{0,}https")
    for i in range(20):
        print(document.get_document().paragraphs[lists[i]].text)
